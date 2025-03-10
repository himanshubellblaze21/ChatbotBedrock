# import boto3
# import os
# import fitz  # PyMuPDF for PDF extraction
# import docx

# s3_client = boto3.client("s3")

# def extract_text(file_name, file_type, s3_bucket):
#     """Download file from S3 and extract text from a PDF or DOCX file."""

#     # Ensure a correct temporary directory
#     temp_dir = os.path.join(os.getcwd(), "temp_files")
#     os.makedirs(temp_dir, exist_ok=True)  # Create temp folder if not exists
#     local_path = os.path.join(temp_dir, file_name)

#     # Download file from S3
#     s3_client.download_file(s3_bucket, file_name, local_path)

#     text = ""

#     try:
#         if file_type == "pdf":
#             # Use `with` statement to ensure proper closure of the file
#             with fitz.open(local_path) as doc:
#                 text = "\n".join([page.get_text("text") for page in doc])

#         elif file_type == "docx":
#             # Read DOCX file properly
#             doc = docx.Document(local_path)
#             text = "\n".join([para.text for para in doc.paragraphs])

#     except Exception as e:
#         print(f"❌ Error extracting text: {e}")

#     finally:
#         # Ensure the file is closed before deleting it
#         if os.path.exists(local_path):
#             try:
#                 os.remove(local_path)
#             except PermissionError:
#                 print(f"⚠️ Could not delete {local_path}, it may still be in use.")

#     return text
import boto3
import os
import fitz  # PyMuPDF for PDF text extraction
import docx
import pytesseract
from pdf2image import convert_from_path

# Initialize S3 client
s3_client = boto3.client("s3")

# Function to extract text from a PDF file (including OCR)
def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file. Uses OCR if no text is found."""
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text("text") for page in doc])

    # If no text is found, use OCR (for scanned PDFs)
    if not text.strip():
        print("⚠️ No text found in PDF, attempting OCR...")
        images = convert_from_path(pdf_path)
        for img in images:
            text += pytesseract.image_to_string(img)

    return text.strip() if text.strip() else "Error: No extractable text found."

# Function to extract text from a DOCX file
def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip() if text.strip() else "Error: No extractable text found."
    except Exception as e:
        return f"Error extracting DOCX text: {e}"

# Function to download file from S3 and extract text
def extract_text(file_name, file_type, s3_bucket):
    """Download file from S3 and extract text from a PDF or DOCX file."""
    
    # Ensure a correct temporary directory
    temp_dir = os.path.join(os.getcwd(), "temp_files")
    os.makedirs(temp_dir, exist_ok=True)  # Create temp folder if not exists
    local_path = os.path.join(temp_dir, file_name)

    # Download file from S3
    try:
        s3_client.download_file(s3_bucket, file_name, local_path)
    except Exception as e:
        return f"❌ Error downloading file from S3: {e}"

    text = ""

    try:
        if file_type.lower() == "pdf":
            text = extract_text_from_pdf(local_path)
        elif file_type.lower() == "docx":
            text = extract_text_from_docx(local_path)
        else:
            text = "❌ Unsupported file type. Please upload a PDF or DOCX."

    except Exception as e:
        text = f"❌ Error extracting text: {e}"

    finally:
        # Ensure the file is closed before deleting it
        if os.path.exists(local_path):
            try:
                os.remove(local_path)
            except PermissionError:
                print(f"⚠️ Could not delete {local_path}, it may still be in use.")

    return text
